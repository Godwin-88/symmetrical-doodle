"""RAG (Retrieval Augmented Generation) service for financial research."""

import asyncio
import os
import hashlib
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime, timezone
import aiofiles
import chromadb
from chromadb.config import Settings
import pandas as pd
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
from bs4 import BeautifulSoup
import json

from .llm_config import LLMIntelligenceConfig, load_llm_config
from .llm_service import LLMService
from .logging import get_logger

logger = get_logger(__name__)


class Document:
    """Document model for RAG system."""
    
    def __init__(
        self,
        content: str,
        metadata: Dict[str, Any],
        doc_id: str = None,
        embedding: List[float] = None
    ):
        self.content = content
        self.metadata = metadata
        self.doc_id = doc_id or self._generate_id()
        self.embedding = embedding
        self.created_at = datetime.now(timezone.utc)
    
    def _generate_id(self) -> str:
        """Generate unique document ID."""
        content_hash = hashlib.md5(self.content.encode()).hexdigest()
        return f"doc_{content_hash[:12]}"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "doc_id": self.doc_id,
            "content": self.content,
            "metadata": self.metadata,
            "created_at": self.created_at.isoformat(),
            "embedding": self.embedding
        }


class DocumentProcessor:
    """Document processing and chunking."""
    
    def __init__(self, config: LLMIntelligenceConfig):
        self.config = config.rag
    
    async def process_file(self, file_path: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """Process file and extract documents."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        metadata = metadata or {}
        metadata.update({
            "file_path": file_path,
            "file_type": file_ext,
            "processed_at": datetime.now(timezone.utc).isoformat()
        })
        
        if file_ext == ".pdf":
            return await self._process_pdf(file_path, metadata)
        elif file_ext == ".txt":
            return await self._process_text(file_path, metadata)
        elif file_ext == ".md":
            return await self._process_markdown(file_path, metadata)
        elif file_ext == ".csv":
            return await self._process_csv(file_path, metadata)
        elif file_ext == ".json":
            return await self._process_json(file_path, metadata)
        elif file_ext == ".html":
            return await self._process_html(file_path, metadata)
        elif file_ext in [".doc", ".docx"]:
            return await self._process_docx(file_path, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    async def _process_pdf(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process PDF file."""
        documents = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            full_text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                full_text += page_text + "\n"
                
                # Create page-level documents for large PDFs
                if len(page_text.strip()) > 100:
                    page_metadata = metadata.copy()
                    page_metadata.update({"page_number": page_num + 1})
                    
                    chunks = self._chunk_text(page_text)
                    for i, chunk in enumerate(chunks):
                        chunk_metadata = page_metadata.copy()
                        chunk_metadata.update({"chunk_index": i})
                        documents.append(Document(chunk, chunk_metadata))
        
        return documents
    
    async def _process_text(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process text file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
        
        chunks = self._chunk_text(content)
        documents = []
        
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update({"chunk_index": i})
            documents.append(Document(chunk, chunk_metadata))
        
        return documents
    
    async def _process_csv(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process CSV file."""
        df = pd.read_csv(file_path)
        documents = []
        
        # Create document from CSV summary
        summary = f"CSV Data Summary:\n"
        summary += f"Columns: {', '.join(df.columns)}\n"
        summary += f"Rows: {len(df)}\n"
        summary += f"Data types:\n{df.dtypes.to_string()}\n"
        summary += f"Statistical summary:\n{df.describe().to_string()}"
        
        metadata.update({"data_type": "csv_summary"})
        documents.append(Document(summary, metadata))
        
        # Create documents from individual rows (for small datasets)
        if len(df) <= 1000:
            for idx, row in df.iterrows():
                row_text = f"Row {idx}:\n" + "\n".join([f"{col}: {val}" for col, val in row.items()])
                row_metadata = metadata.copy()
                row_metadata.update({"row_index": idx, "data_type": "csv_row"})
                documents.append(Document(row_text, row_metadata))
        
        return documents
    
    async def _process_json(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process JSON file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
            data = json.loads(content)
        
        # Convert JSON to readable text
        json_text = json.dumps(data, indent=2)
        chunks = self._chunk_text(json_text)
        documents = []
        
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update({"chunk_index": i, "data_type": "json"})
            documents.append(Document(chunk, chunk_metadata))
        
        return documents
    
    async def _process_html(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process HTML file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
        
        soup = BeautifulSoup(content, 'html.parser')
        text_content = soup.get_text()
        
        chunks = self._chunk_text(text_content)
        documents = []
        
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update({"chunk_index": i, "data_type": "html"})
            documents.append(Document(chunk, chunk_metadata))
        
        return documents
    
    def _chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces."""
        if len(text) <= self.config.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.config.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for i in range(end, max(start + self.config.chunk_size - self.config.chunk_overlap, start), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
                else:
                    # Look for word boundaries
                    for i in range(end, max(start + self.config.chunk_size - self.config.chunk_overlap, start), -1):
                        if text[i].isspace():
                            end = i
                            break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.config.chunk_overlap
        
        return chunks


class VectorStore:
    """Vector store for document embeddings."""
    
    def __init__(self, config: LLMIntelligenceConfig):
        self.config = config.rag
        self.embedding_model = None
        self.chroma_client = None
        self.collection = None
    
    async def initialize(self):
        """Initialize vector store."""
        # Initialize embedding model
        self.embedding_model = SentenceTransformer(
            self.config.embedding_model,
            device=self.config.embedding_device
        )
        
        # Initialize ChromaDB
        if self.config.vector_db == "chroma":
            self.chroma_client = chromadb.PersistentClient(
                path=self.config.vector_db_path,
                settings=Settings(anonymized_telemetry=False)
            )
            
            self.collection = self.chroma_client.get_or_create_collection(
                name="financial_documents",
                metadata={"hnsw:space": "cosine"}
            )
        
        logger.info("Vector store initialized")
    
    async def add_documents(self, documents: List[Document]):
        """Add documents to vector store."""
        if not documents:
            return
        
        # Generate embeddings
        texts = [doc.content for doc in documents]
        embeddings = self.embedding_model.encode(texts, convert_to_tensor=False)
        
        # Store in ChromaDB
        ids = [doc.doc_id for doc in documents]
        metadatas = [doc.metadata for doc in documents]
        
        self.collection.add(
            ids=ids,
            embeddings=embeddings.tolist(),
            documents=texts,
            metadatas=metadatas
        )
        
        logger.info(f"Added {len(documents)} documents to vector store")
    
    async def search(self, query: str, top_k: int = None) -> List[Tuple[Document, float]]:
        """Search for similar documents."""
        top_k = top_k or self.config.top_k_retrieval
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode([query], convert_to_tensor=False)[0]
        
        # Search in ChromaDB
        results = self.collection.query(
            query_embeddings=[query_embedding.tolist()],
            n_results=top_k,
            include=["documents", "metadatas", "distances"]
        )
        
        # Convert results to Document objects
        documents = []
        for i in range(len(results['ids'][0])):
            doc = Document(
                content=results['documents'][0][i],
                metadata=results['metadatas'][0][i],
                doc_id=results['ids'][0][i]
            )
            similarity = 1 - results['distances'][0][i]  # Convert distance to similarity
            
            if similarity >= self.config.similarity_threshold:
                documents.append((doc, similarity))
        
        return documents


class RAGService:
    """Main RAG service orchestrator."""
    
    def __init__(self, config: LLMIntelligenceConfig = None):
        self.config = config or load_llm_config()
        self.document_processor = DocumentProcessor(self.config)
        self.vector_store = VectorStore(self.config)
        self.llm_service = LLMService(self.config)
    
    async def initialize(self):
        """Initialize RAG service."""
        await self.vector_store.initialize()
        await self.llm_service.initialize()
        logger.info("RAG Service initialized")
    
    async def ingest_document(self, file_path: str, metadata: Dict[str, Any] = None) -> int:
        """Ingest document into RAG system."""
        try:
            documents = await self.document_processor.process_file(file_path, metadata)
            await self.vector_store.add_documents(documents)
            
            logger.info(f"Ingested {len(documents)} chunks from {file_path}")
            return len(documents)
            
        except Exception as e:
            logger.error(f"Failed to ingest document {file_path}: {e}")
            raise
    
    async def query(self, question: str, context_filter: Dict[str, Any] = None) -> Dict[str, Any]:
        """Query RAG system with question."""
        try:
            # Retrieve relevant documents
            relevant_docs = await self.vector_store.search(question)
            
            if not relevant_docs:
                # No relevant documents found, use LLM without context
                response = await self.llm_service.financial_analysis(question)
                return {
                    "answer": response.content,
                    "sources": [],
                    "confidence": "low",
                    "model_info": {
                        "model": response.model,
                        "provider": response.provider,
                        "tokens_used": response.tokens_used,
                        "response_time": response.response_time
                    }
                }
            
            # Prepare context from retrieved documents
            context_docs = []
            sources = []
            
            for doc, similarity in relevant_docs[:5]:  # Use top 5 documents
                context_docs.append(f"Source: {doc.metadata.get('file_path', 'Unknown')}\n{doc.content}")
                sources.append({
                    "file_path": doc.metadata.get('file_path', 'Unknown'),
                    "similarity": similarity,
                    "metadata": doc.metadata
                })
            
            # Generate answer with context
            response = await self.llm_service.research_query(question, context_docs)
            
            # Determine confidence based on similarity scores
            avg_similarity = sum(sim for _, sim in relevant_docs[:3]) / min(3, len(relevant_docs))
            confidence = "high" if avg_similarity > 0.8 else "medium" if avg_similarity > 0.6 else "low"
            
            return {
                "answer": response.content,
                "sources": sources,
                "confidence": confidence,
                "model_info": {
                    "model": response.model,
                    "provider": response.provider,
                    "tokens_used": response.tokens_used,
                    "response_time": response.response_time
                },
                "retrieval_info": {
                    "documents_found": len(relevant_docs),
                    "avg_similarity": avg_similarity
                }
            }
            
        except Exception as e:
            logger.error(f"RAG query failed: {e}")
            raise
    
    async def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about ingested documents."""
        try:
            collection_info = self.vector_store.collection.count()
            
            return {
                "total_documents": collection_info,
                "vector_store": self.config.rag.vector_db,
                "embedding_model": self.config.rag.embedding_model,
                "chunk_size": self.config.rag.chunk_size,
                "similarity_threshold": self.config.rag.similarity_threshold
            }
            
        except Exception as e:
            logger.error(f"Failed to get document stats: {e}")
            return {"error": str(e)}